import os
import requests
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pinecone.grpc import PineconeGRPC as Pinecone
from docxtpl import DocxTemplate
import string
import secrets
import json
import boto3

from dotenv import load_dotenv
# .env 파일 로드
load_dotenv()

TEMP_DOCX_DIR = os.getenv("TEMP_DOCX_DIR")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
AWS_ACCESS_KEY = os.getenv("AWS_ACCESS_KEY")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
BUCKET_NAME= os.getenv("BUCKET_NAME")
DOCX_OUTPUT_DIR= os.getenv("DOCX_OUTPUT_DIR")

s3 = boto3.client(
    "s3",
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
)

def download_docx_from_s3(url: str) -> str:
    """
    주어진 S3 URL에서 .docx 파일을 다운로드하여
    OS 임시 디렉토리에 저장한 뒤 로컬 경로를 반환합니다.
    """
    # 1) HTTP GET 요청으로 파일 가져오기
    resp = requests.get(url)                    # requests 사용 :contentReference[oaicite:1]{index=1}
    resp.raise_for_status()

    # 2) 임시 디렉토리 및 파일 경로 생성
    temp_dir = TEMP_DOCX_DIR
    filename = os.path.basename(url)
    local_path = os.path.join(temp_dir, filename)

    # 3) 바이너리 쓰기
    with open(local_path, "wb") as f:
        f.write(resp.content)

    return local_path, filename

def extract_docx_text(path: str) -> dict:
    """
    .docx 파일 경로를 받아, 본문(paragraphs), 표(tables),
    헤더(headers), 푸터(footers) 텍스트를 추출하여 반환합니다.
    """
    doc = Document(path)
    data = {"paragraphs": [], "tables": [], "headers": [], "footers": []}

    # 본문: 문서 순서대로 paragraph와 table 처리
    for block in doc.iter_inner_content():
        if isinstance(block, Paragraph):
            data["paragraphs"].append(block.text)
        elif isinstance(block, Table):
            table_data = [
                [cell.text for cell in row.cells]
                for row in block.rows
            ]
            data["tables"].append(table_data)

    # 각 section의 header/footer
    for section in doc.sections:
        for para in section.header.paragraphs:
            data["headers"].append(para.text)
        for para in section.footer.paragraphs:
            data["footers"].append(para.text)

    return data

def index_pdfs(extract_doc, filename):
    """
    For each PDF path, extract chunks, compute embeddings, and upsert to Pinecone.
    """
    pinecone = Pinecone(
        api_key=PINECONE_API_KEY
        )
    INDEX_NAME = "agile"
    NAMESPACE = "ns1"
    index = pinecone.Index(INDEX_NAME)

    embeddings = pinecone.inference.embed(
        model="multilingual-e5-large",
        inputs=str(extract_doc),
        parameters={"input_type": "passage", "truncate": "END"}
    )
    print(embeddings.get('data'))
    vec_id = filename
    metadata = {"source": filename, "text": str(extract_doc)}
    vectors.append((vec_id, embeddings.get("data")[0]['values'], metadata))
    # Batch upsert every 100 vectors
    if len(vectors) >= 100:
        index.upsert(vectors=vectors, namespace=NAMESPACE)  
        vectors = []
    # Upsert any remaining vectors
    if vectors:
        index.upsert(vectors=vectors, namespace=NAMESPACE)

def docx_to_pine(url: str):
    """
    S3 URL에서 .docx 파일을 다운로드하고,
    텍스트를 추출하여 Pinecone에 인덱싱합니다.
    """
    local_path, filename = download_docx_from_s3(url)
    extract_doc = extract_docx_text(local_path)
    index_pdfs(extract_doc, filename)
    return True

def random_key(length: int = 20) -> str:
    alphabet = string.ascii_letters + string.digits
    return ''.join(secrets.choice(alphabet) for _ in range(length))

def docx_to_s3(template_content: dict, templateId: str):
    """
    주어진 minutes 딕셔너리를 사용하여
    템플릿 .docx 파일을 생성하고 S3에 업로드합니다.
    """
    
    # 템플릿 워드 파일 불러오기
    if templateId == '1':
        doc = DocxTemplate("template/template1.docx")
        file_name = f"template1_{random_key()}.docx"
    elif templateId == '2':
        doc = DocxTemplate("template/template2.docx")
        file_name = f"template2_{random_key()}.docx"
    elif templateId == '3':
        doc = DocxTemplate("template/template3.docx")
        file_name = f"template3_{random_key()}.docx"

    context = {item.objectKey.replace(' ', '_'): item.objectValue for item in template_content}

    # 템플릿에 값 채워넣기
    doc.render(context)

    file_path = f"{DOCX_OUTPUT_DIR}/{file_name}"
    doc.save(file_path)

    s3_key = f"docx/{file_name}"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    s3.upload_fileobj(
            Fileobj=open(file_path, "rb"),
            Bucket=BUCKET_NAME,
            Key=s3_key,
            ExtraArgs={"ContentType": content_type}
        )
    
    file_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{s3_key}"

    return file_url   
